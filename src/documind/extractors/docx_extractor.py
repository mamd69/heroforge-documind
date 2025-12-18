"""
Word Document (.docx) Extraction
Handles paragraphs, tables, and metadata
"""
from docx import Document
from typing import Dict, List
from pathlib import Path

class DocxExtractor:
    """Extract content from Word documents"""

    def extract(self, docx_path: str) -> Dict:
        """
        Extract all content from a Word document.

        Args:
            docx_path: Path to .docx file

        Returns:
            Dictionary with text, tables, and metadata
        """
        try:
            doc = Document(docx_path)

            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
            }

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name
                    })

            # Extract tables
            tables = []
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    "table_number": table_num + 1,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                    "data": table_data
                })

            # Combine all text
            full_text = "\n\n".join(para["text"] for para in paragraphs)

            return {
                "success": True,
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to extract from {docx_path}"
            }

    def extract_text_only(self, docx_path: str) -> str:
        """Extract just the text content"""
        result = self.extract(docx_path)
        return result.get("text", "") if result["success"] else ""

    def format_for_llm(self, docx_path: str) -> str:
        """
        Format document content for LLM consumption.
        Combines text and tables in a readable format.
        """
        result = self.extract(docx_path)

        if not result["success"]:
            return ""

        output = []

        # Add metadata as header
        if result["metadata"]["title"]:
            output.append(f"# {result['metadata']['title']}\n")

        # Add paragraphs
        for para in result["paragraphs"]:
            # Check if it's a heading
            if para["style"].startswith("Heading"):
                level = para["style"][-1]
                output.append(f"{'#' * int(level)} {para['text']}\n")
            else:
                output.append(f"{para['text']}\n")

        # Add tables in markdown format
        for table in result["tables"]:
            output.append(f"\n## Table {table['table_number']}\n")

            data = table["data"]
            if not data:
                continue

            # Header
            output.append("| " + " | ".join(data[0]) + " |")
            output.append("|" + "|".join(["---"] * len(data[0])) + "|")

            # Rows
            for row in data[1:]:
                output.append("| " + " | ".join(row) + " |")

            output.append("")

        return "\n".join(output)

# Test
if __name__ == "__main__":
    # Test extraction using the sample document
    extractor = DocxExtractor()
    result = extractor.extract("docs/workshops/S7-sample-docs/meeting_notes.docx")

    print("=" * 60)
    print("DOCX EXTRACTION TEST")
    print("=" * 60)
    print(f"Title: {result['metadata']['title']}")
    print(f"Author: {result['metadata']['author']}")
    print(f"Paragraphs: {result['paragraph_count']}")
    print(f"Tables: {result['table_count']}")

    print("\n" + "-" * 60)
    print("LLM-Formatted Output:")
    print("-" * 60)
    print(extractor.format_for_llm("docs/workshops/S7-sample-docs/meeting_notes.docx"))